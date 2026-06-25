import os
import uuid
from datetime import datetime
from typing import Dict, Optional, Tuple

from docx import Document


def _ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)


def _safe_file_name(value: str) -> str:
    keep = []
    for char in value:
        keep.append(char if char.isalnum() or char in {"-", "_", "."} else "_")
    return "".join(keep)


def _render_resume_text(profile: Dict, cv_result: Dict, job: Dict) -> str:
    lines = []
    name = profile.get("name") or "Candidate"
    lines.append(name)
    lines.append("")
    if profile.get("email"):
        lines.append(f"Email: {profile.get('email')}")
    if profile.get("phone"):
        lines.append(f"Phone: {profile.get('phone')}")
    if profile.get("target_location"):
        lines.append(f"Location: {profile.get('target_location')}")
    lines.append("")

    lines.append("Professional Summary")
    lines.append(
        f"Role-focused profile aligned to {job.get('title', 'target role')} opportunities with practical experience in "
        f"{', '.join((profile.get('skills') or [])[:6]) or 'software development'}."
    )
    lines.append("")

    lines.append("Key Skills")
    for skill in (profile.get("skills") or [])[:20]:
        lines.append(f"- {skill}")
    lines.append("")

    lines.append("Experience")
    for exp in profile.get("experiences", []) or []:
        role = exp.get("role", "Role")
        company = exp.get("company", "Company")
        period = f"{exp.get('start_year', '')} - {exp.get('end_year', '')}".strip()
        lines.append(f"{role} | {company} | {period}")
        if exp.get("description"):
            lines.append(f"  {exp.get('description')}")
    lines.append("")

    lines.append("Education")
    for edu in profile.get("education", []) or []:
        degree = edu.get("degree", "Degree")
        school = edu.get("university", "University")
        year = edu.get("year")
        year_str = f" ({year})" if year else ""
        lines.append(f"- {degree}, {school}{year_str}")
    lines.append("")

    if cv_result.get("keywords_added"):
        lines.append("Role-Targeted Keywords Integrated")
        for keyword in cv_result.get("keywords_added", [])[:15]:
            lines.append(f"- {keyword}")

    return "\n".join(lines).strip()


def _render_cover_letter_text(profile: Dict, cover_letter: str) -> str:
    user_name = (profile.get("name") or "").strip() or "Candidate"
    if cover_letter.strip().lower().endswith("clover"):
        cover_letter = cover_letter.rsplit("Clover", 1)[0].rstrip()
    if "[Your Name]" in cover_letter:
        cover_letter = cover_letter.replace("[Your Name]", user_name)
    if user_name not in cover_letter[-120:]:
        cover_letter = f"{cover_letter.strip()}\n\nSincerely,\n{user_name}\n"
    return cover_letter


def _write_docx(path: str, title: str, body: str) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for block in body.split("\n\n"):
        if not block.strip():
            continue
        if "\n" in block and all(line.strip().startswith("- ") for line in block.split("\n")[1:]):
            lines = block.split("\n")
            doc.add_heading(lines[0], level=2)
            for line in lines[1:]:
                doc.add_paragraph(line.replace("- ", "", 1), style="List Bullet")
        else:
            doc.add_paragraph(block)
    doc.save(path)


def create_generation_assets(
    output_root: str,
    user_id: Optional[str],
    profile: Dict,
    job: Dict,
    cv_result: Dict,
    cover_letter: str,
) -> Tuple[str, str]:
    owner = user_id or "anonymous"
    folder = os.path.join(
        output_root,
        "generated_assets",
        _safe_file_name(owner),
        _safe_file_name(profile.get("id", "profile")),
        datetime.utcnow().strftime("%Y%m%d"),
    )
    _ensure_dir(folder)

    resume_body = _render_resume_text(profile=profile, cv_result=cv_result, job=job)
    cover_text = _render_cover_letter_text(profile=profile, cover_letter=cover_letter)

    token = uuid.uuid4().hex[:10]
    resume_file = os.path.join(folder, f"optimized_resume_{token}.docx")
    cover_file = os.path.join(folder, f"cover_letter_{token}.docx")

    _write_docx(path=resume_file, title="Optimized Resume", body=resume_body)
    _write_docx(path=cover_file, title="Cover Letter", body=cover_text)
    return resume_file, cover_file
